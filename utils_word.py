from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("INDUSTRIA SIGRAMA S.A. DE C.V.")
    run.bold = True
    run.font.size = Pt(11)
    
    # Text right aligned
    run2 = paragraph.add_run("\t\t\tRevisión 01")
    run2.font.size = Pt(8)

def crear_word_generico_muestra(titulo, columnas, filas, output_path):
    doc = Document()
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(10)

    add_header_footer(doc)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"MUESTRA DE FORMATO OFICIAL: {titulo}")
    run_title.bold = True
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = RGBColor(0xEC, 0x20, 0x24) # Corporate Red #EC2024
    
    doc.add_paragraph("Este documento es una muestra de control regulada bajo el Sistema de Gestión de Calidad (SGC) de Industria SIGRAMA.")
    
    # Table
    if columnas and filas:
        table = doc.add_table(rows=1, cols=len(columnas))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columnas):
            hdr_cells[i].text = str(col_name)
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            
        for fila in filas:
            row_cells = table.add_row().cells
            for i, cell_val in enumerate(fila):
                row_cells[i].text = str(cell_val)
                
    doc.add_paragraph("\n")
    
    # Signatures table
    sig_table = doc.add_table(rows=1, cols=2)
    sig_cells = sig_table.rows[0].cells
    sig_cells[0].text = "_____________________________\nELABORÓ (DEPARTAMENTO DE CALIDAD)"
    sig_cells[1].text = "_____________________________\nAPROBÓ (CONTROL DOCUMENTAL / SGC)"
    sig_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_path)
